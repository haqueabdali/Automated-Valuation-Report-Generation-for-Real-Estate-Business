from data_processor import DataProcessor
from valuation_calculator import ValuationCalculator
from report_generator import ReportGenerator
from src.config import DEFAULT_METHOD, VALUATION_METHODS
from docx import Document
from pathlib import Path
#from data_processing.data_loader import load_data
from src.data_processing.data_loader import load_data
from valuation_model.model import predict_valuation
from report_generation.report_generator import generate_report
import matplotlib
import pandas as pd
import argparse 
matplotlib.use('Agg')
# Add project root to Python path
sys.path.append(str(Path(__file__).parent))

from data_processing.data_loader import load_data 
import sys
sys.path.append(str(Path(__file__).parent))

def parse_arguments():
    """Improved argument parsing with custom error handling"""
    parser = argparse.ArgumentParser(
        description="Real Estate Valuation Report Generator",
        add_help=False  # We'll add custom help handling
    )
    
    parser = argparse.ArgumentParser(description='Property Valuation Report Generator')
    parser.add_argument('--csv', type=str, help='Path to CSV input file')
    parser.add_argument('--urls', nargs='+', help='Property URLs to scrape')
    parser.add_argument('--property-id', type=int, help='Property ID for database lookup')
    parser.add_argument('--method', choices=VALUATION_METHODS, default='hybrid',
                      help='Valuation method to use')
    parser.add_argument('--all-methods', action='store_true',
                      help='Run all valuation methods')
    try:
        args = parser.parse_args()
        return args
    except argparse.ArgumentError as e:
        print(f"Error: {str(e)}")
        parser.print_help()
        sys.exit(2)
    except Exception as e:
        print(f"Unexpected error: {str(e)}")
        sys.exit(1)
        
    return parser.parse_args()

    # # Required arguments
    # required = parser.add_argument_group('required arguments')
    # required.add_argument(
    #     'property_id',
    #     type=str,
    #     help="ID of the property to evaluate"
    # )
    
    # # Optional arguments
    # optional = parser.add_argument_group('optional arguments')
    # optional.add_argument(
    #     '-m', '--method',
    #     type=str,
    #     default=DEFAULT_METHOD,
    #     choices=VALUATION_METHODS,
    #     help=f"Valuation method to use (default: {DEFAULT_METHOD})"
    # )
    # optional.add_argument(
    #     '-a', '--all-methods',
    #     action='store_true',
    #     help="Run all valuation methods and include in report"
    # )
    # optional.add_argument(
    #     '-h', '--help',
    #     action='help',
    #     help="Show this help message and exit"
    # )
    
    # Custom error handling

        
# Create templates directory if it doesn't exist
template_dir = Path("templates")
template_dir.mkdir(exist_ok=True)

# Create a brand new valid Word document
doc = Document()

# Add basic structure with placeholders
doc.add_heading("VALUATION REPORT", level=1)
doc.add_paragraph("Company: {company_name}")
doc.add_paragraph("Date: {report_date}")
doc.add_heading("PROPERTY DETAILS", level=2)
doc.add_paragraph("Address: {property.address}")
doc.add_paragraph("City: {property.city}")
doc.add_paragraph("State: {property.state}")
doc.add_paragraph("Type: {property.property_type}")

# Save the document
template_path = template_dir / "report_template.docx"
doc.save(template_path)
print(f"Created new valid template at: {template_path}")

def main():
    args = parse_arguments()
    
    try:
        # Initialize components
        data_processor = DataProcessor()
        valuation_calculator = ValuationCalculator(data_processor)
        report_generator = ReportGenerator()
        
        # Load and validate data - MODIFIED SECTION
        if args.csv:
            print(f"Loading data from CSV: {args.csv}")
            if not data_processor.load_data_from_csv(args.csv):
                print("Failed to load CSV data")
                return
        elif args.urls:
            print(f"Scraping data from URLs: {args.urls}")
            scraped_data = data_processor.scrape_property_data(args.urls)
            if not scraped_data:
                print("Failed to scrape property data")
                return
        else:
            if not data_processor.load_data():
                print("Failed to load default data")
                return
    
        # Get property details
        if args.urls:
            # Use first scraped property if URLs provided
            property_details = data_processor.get_scraped_property_details()
        
        else:
            property_details = data_processor.get_property_details(args.property_id)
        
        if not args.urls:
            error_msg += f" with ID {args.property_id}"    
        
        if not property_details:
            error_msg = f"No property found"
            if not args.urls:
                error_msg += f" with ID {args.property_id}"
            print(error_msg)
            sys.exit(1)
        
        # Calculate valuation
        if args.all_methods:
            print("Running all valuation methods...")
            valuation_results = []
            for method in VALUATION_METHODS:
                print(f"Calculating {method}...")
                result = valuation_calculator.calculate_valuation(property_details, method)
                if result:
                    valuation_results.append(result)
            
            # Sort by confidence (highest first)
            valuation_results.sort(key=lambda x: x.confidence, reverse=True)
        else:
            print(f"Calculating valuation using {args.method} method...")
            valuation_results = valuation_calculator.calculate_valuation(property_details, args.method)
    
        # Generate report
        report_path = report_generator.generate_report(
            property_details, 
            valuation_results,
            data_source="Web Scraping" if args.urls else ("CSV" if args.csv else "Database")
        )
        print(f"\nSuccessfully generated valuation report:\n{report_path}")
    
    except Exception as e:
        print(f"\nError during report generation: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()