from docxtpl import DocxTemplate
from src.config import REPORT_TEMPLATE, OUTPUT_DIR, COMPANY_NAME
from datetime import datetime
from pathlib import Path
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt
import logging
import os

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ReportGenerator:
    DEFAULT_PROPERTY_VALUES = {
        'address': 'N/A',
        'city': 'N/A',
        'state': 'N/A',
        'zip_code': 'N/A',
        'property_type': 'N/A',
        'bedrooms': 'N/A',
        'bathrooms': 'N/A',
        'sqft': 'N/A',
        'lot_size': 'N/A',
        'year_built': 'N/A'
    }

    def __init__(self, template_path: str = None, company_name: str = None):
        self.template_path = Path(template_path) if template_path else Path(REPORT_TEMPLATE)
        self.output_dir = Path(OUTPUT_DIR)
        self.company_name = company_name or COMPANY_NAME
        self.ensure_directories_exist()

    def ensure_directories_exist(self):
        """Ensure required directories exist"""
        self.template_path.parent.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def create_default_template(self):
        """Generate a professional default template"""
        doc = Document()
        
        # Set document styles
        styles = doc.styles
        title_style = styles['Title']
        title_style.font.size = Pt(20)
        title_style.font.bold = True
        
        heading1 = styles['Heading 1']
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        
        # Document header
        doc.add_paragraph(self.company_name, style='Title')
        doc.add_paragraph('PROPERTY VALUATION REPORT', style='Heading 1')
        doc.add_paragraph('Date: {{report_date}}')
        doc.add_paragraph('Report ID: {{report_id}}')
        
        # Property details section
        doc.add_paragraph('PROPERTY INFORMATION', style='Heading 1')
        doc.add_paragraph('Address: {{property.address}}')
        doc.add_paragraph('City: {{property.city}}')
        doc.add_paragraph('State: {{property.state}} | ZIP: {{property.zip_code}}')
        doc.add_paragraph('Type: {{property.property_type}}')
        doc.add_paragraph('Bedrooms: {{property.bedrooms}} | Bathrooms: {{property.bathrooms}}')
        doc.add_paragraph('Sq Ft: {{property.sqft}} | Lot Size: {{property.lot_size}}')
        doc.add_paragraph('Year Built: {{property.year_built}}')
        
        # Valuation results section
        doc.add_paragraph('VALUATION ANALYSIS', style='Heading 1')
        doc.add_paragraph('Primary Method: {{valuation.primary_method}}')
        doc.add_paragraph('Estimated Value: {{valuation.primary_value}}')
        doc.add_paragraph('Confidence Level: {{valuation.primary_confidence}}')
        
        # Comparative analysis section
        doc.add_paragraph('COMPARATIVE ANALYSIS', style='Heading 1')
        doc.add_paragraph('{{comparative_chart}}')
        
        # Save template
        doc.save(self.template_path)
        logger.info(f"Created new template at {self.template_path}")

    def ensure_valid_template(self):
        """Ensure a valid template exists"""
        if not self.template_path.exists() or self.template_path.stat().st_size < 1024:
            logger.warning("Template missing or invalid - creating default")
            self.create_default_template()
        return self.template_path

    def generate_report(self, property_details: dict, valuation_results: list, data_source: str = None) -> Path:
        """Generate a professional valuation report"""
        try:
            # Prepare context data
            context = self.prepare_context(property_details, valuation_results)
            
            # Generate charts
            context['comparative_chart'] = self.generate_comparison_chart(
                valuation_results, 
                context['report_id']
            )
            
            # Render and save report
            doc = DocxTemplate(self.ensure_valid_template())
            doc.render(context)
            
            output_path = self.output_dir / f"{context['report_id']}.docx"
            doc.save(output_path)
            
            logger.info(f"Report generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Report generation failed: {str(e)}")
            # Create minimal error report
            error_path = self.output_dir / f"error_report_{datetime.now().strftime('%H%M%S')}.txt"
            with open(error_path, 'w') as f:
                f.write(f"Report Generation Error\n{str(e)}")
            return error_path

    def prepare_context(self, property_details: dict, valuation_results: list) -> dict:
        """Prepare data context for report rendering"""
        # Generate unique report ID
        report_id = f"REP-{property_details.get('id', 'UNK')}-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
        
        # Format valuation data
        valuation_data = self.format_valuation_data(valuation_results)
        
        return {
            'company_name': self.company_name,
            'report_date': datetime.now().strftime("%B %d, %Y"),
            'report_id': report_id,
            'property': {**self.DEFAULT_PROPERTY_VALUES, **property_details},
            'valuation': valuation_data
        }

    def format_valuation_data(self, valuation_results: list) -> dict:
        """Format valuation results for reporting"""
        primary = valuation_results[0] if isinstance(valuation_results, list) else valuation_results
        
        return {
            'primary_method': primary.method.replace('_', ' ').title(),
            'primary_value': self.format_currency(primary.value),
            'primary_confidence': f"{int(primary.confidence * 100)}%",
            'final_value': self.format_currency(primary.value)
        }

    def format_currency(self, value: float) -> str:
        """Format currency values consistently"""
        return "${:,.2f}".format(value) if isinstance(value, (int, float)) else value

    def generate_comparison_chart(self, valuation_results: list, report_id: str) -> str:
        """Generate professional comparative analysis chart"""
        try:
            if not isinstance(valuation_results, list) or not valuation_results:
                return None
                
            primary = valuation_results[0]
            if primary.method != "sales_comparison":
                return None
                
            # Extract comparable data
            comp_data = primary.details.get('comparables', [])
            adjustments = primary.details.get('adjustments', [])
            
            if not comp_data or not adjustments:
                return None
                
            # Prepare chart data
            comp_ids = [f"Comp {i+1}" for i in range(len(comp_data))]
            sale_prices = [c.get('sale_price', 0) for c in comp_data]
            adjusted_prices = [a.get('adjusted_price', 0) for a in adjustments]
            
            # Create professional chart
            plt.style.use('seaborn-whitegrid')
            fig, ax = plt.subplots(figsize=(10, 6))
            
            x = range(len(comp_ids))
            bar_width = 0.35
            
            ax.bar(x, sale_prices, bar_width, label='Original Price', color='#1f77b4')
            ax.bar([i + bar_width for i in x], adjusted_prices, bar_width, 
                    label='Adjusted Price', color='#ff7f0e')
            
            # Add property line
            ax.axhline(y=primary.value, color='r', linestyle ='--', 
                        label='Subject Property Value')
            
            # Format chart
            ax.set_xlabel('Comparable Properties')
            ax.set_ylabel('Value (USD)')
            ax.set_title('Sales Comparison Analysis')
            ax.set_xticks([i + bar_width/2 for i in x])
            ax.set_xticklabels(comp_ids)
            ax.legend()
            ax.grid(True, linestyle='--', alpha=0.7)
            fig.tight_layout()
            
            # Save chart
            chart_path = self.output_dir / f"comparison_{report_id}.png"
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            plt.close()
            
            return str(chart_path)
            
        except Exception as e:
            logger.error(f"Chart generation failed: {str(e)}")
            return None