from docxtpl import DocxTemplate
from src.config import REPORT_TEMPLATE, OUTPUT_DIR, COMPANY_NAME
from datetime import datetime
from pathlib import Path
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt
import os

class ReportGenerator:
    def __init__(self):
        self.template = None

    def _create_default_template(self):
        """Generate a valid default template if missing"""
        doc = Document()
    
        # Add styles
        styles = doc.styles
        title_style = styles['Title']
        title_style.font.size = Pt(24)
    
        # Add content with placeholders
        doc.add_paragraph('{company_name}', style='Title')
        doc.add_paragraph('VALUATION REPORT', style='Heading 1')
        doc.add_paragraph('{report_date}', style='Heading 2')
    
        # Add property details section
        doc.add_paragraph('PROPERTY DETAILS', style='Heading 1')
        doc.add_paragraph('Address: {property.address}')
        doc.add_paragraph('City: {property.city}')
        doc.add_paragraph('State: {property.state}')
        doc.add_paragraph('ZIP Code: {property.zip_code}')
        doc.add_paragraph('Property Type: {property.property_type}')
        doc.add_paragraph('Bedrooms: {property.bedrooms}')
        doc.add_paragraph('Bathrooms: {property.bathrooms}')
        doc.add_paragraph('Square Footage: {property.sqft}')
        doc.add_paragraph('Lot Size: {property.lot_size}')
        doc.add_paragraph('Year Built: {property.year_built}')
    
        # Add valuation section
        doc.add_paragraph('VALUATION RESULTS', style='Heading 1')
        doc.add_paragraph('Method: {valuation.primary_method}')
        doc.add_paragraph('Value: {valuation.primary_value}')
        doc.add_paragraph('Confidence: {valuation.primary_confidence}')
    
        # Save the document
        template_path = Path(REPORT_TEMPLATE)
        template_path.parent.mkdir(exist_ok=True)
        doc.save(template_path)
        print(f"Created new template at {template_path}")

    def ensure_valid_template(self):
        """Guarantee we have a valid template file"""
        template_path = Path(REPORT_TEMPLATE)
        
        # Create directory if needed
        template_path.parent.mkdir(exist_ok=True)
    
        # Create new template if missing or invalid
        if not template_path.exists() or template_path.stat().st_size == 0:
            print("Creating new valid template...")
            self._create_default_template()
    
        return template_path
        
    def generate_report(self, property_details, valuation_results, output_filename=None):
        """Generate a valuation report with proper template handling"""
        try:
            # Ensure directories exist
            Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
            
            # Prepare context data
            context = {
                'company_name': COMPANY_NAME,
                'report_date': datetime.now().strftime("%B %d, %Y"),
                'property': {
                    'address': property_details.get('address', 'N/A'),
                    'city': property_details.get('city', 'N/A'),
                    'state': property_details.get('state', 'N/A'),
                    'zip_code': property_details.get('zip_code', 'N/A'),
                    'property_type': property_details.get('property_type', 'N/A'),
                    'bedrooms': property_details.get('bedrooms', 'N/A'),
                    'bathrooms': property_details.get('bathrooms', 'N/A'),
                    'sqft': property_details.get('sqft', 'N/A'),
                    'lot_size': property_details.get('lot_size', 'N/A'),
                    'year_built': property_details.get('year_built', 'N/A')
                },
                'valuation': self._prepare_valuation_data(valuation_results),
                'charts': self._generate_charts(valuation_results, property_details.get('id'))
            }
            
            # Ensure template exists
            template_path = self.ensure_valid_template()
            
            # Load template
            doc = DocxTemplate(template_path)
            doc.render(context)
            
            # Save report
            if not output_filename:
                output_filename = f"valuation_{property_details.get('id', 'unknown')}_{datetime.now().strftime('%Y%m%d')}.docx"
            
            output_path = Path(OUTPUT_DIR) / output_filename
            doc.save(output_path)
            
            print(f"Successfully generated report: {output_path}")
            return str(output_path)
            
        except Exception as e:
            print(f"Report generation failed: {str(e)}")
            raise

    def _prepare_valuation_data(self, valuation_results):
        """Prepare valuation data for the report"""
        if not isinstance(valuation_results, list):
            return {
                'primary_method': valuation_results.method.replace('_', ' ').title(),
                'primary_value': "${:,.2f}".format(valuation_results.value),
                'primary_confidence': f"{int(valuation_results.confidence * 100)}%",
                'final_value': "${:,.2f}".format(valuation_results.value)
            }
            
        primary = valuation_results[0]
        return {
            'primary_method': primary.method.replace('_', ' ').title(),
            'primary_value': "${:,.2f}".format(primary.value),
            'primary_confidence': f"{int(primary.confidence * 100)}%",
            'final_value': "${:,.2f}".format(primary.value)
        }
    
    def _generate_charts(self, valuation_results, property_id):
        """Generate charts for the report"""
        charts = {}
        
        try:
            if isinstance(valuation_results, list):
                primary = valuation_results[0]
            else:
                primary = valuation_results
                
            if primary.method == "sales_comparison":
                # Ensure output directory exists
                Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
                
                # Extract data
                comp_data = primary.details.get('comparables', [])
                adjustments = primary.details.get('adjustments', [])
                
                if comp_data and adjustments:
                    comp_ids = [f"Comp {i+1}" for i in range(len(comp_data))]
                    sale_prices = [c.get('sale_price', 0) for c in comp_data]
                    adjusted_prices = [a.get('adjusted_price', 0) for a in adjustments]
                    
                    # Create and save chart
                    plt.figure(figsize=(10, 6))
                    bar_width = 0.35
                    index = range(len(comp_ids))
                    
                    plt.bar(index, sale_prices, bar_width, label='Original Price')
                    plt.bar([i + bar_width for i in index], adjusted_prices, bar_width, label='Adjusted Price')
                    
                    plt.xlabel('Comparable Properties')
                    plt.ylabel('Price ($)')
                    plt.title('Sales Comparison Analysis')
                    plt.xticks([i + bar_width/2 for i in index], comp_ids)
                    plt.legend()
                    plt.tight_layout()
                    
                    chart_path = Path(OUTPUT_DIR) / f"sales_comparison_{property_id}.png"
                    plt.savefig(chart_path, dpi=300)
                    plt.close()
                    
                    charts['sales_comparison'] = str(chart_path)
                    
        except Exception as e:
            print(f"Chart generation error: {str(e)}")
        
        return charts